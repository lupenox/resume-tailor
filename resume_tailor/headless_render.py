from __future__ import annotations

from pathlib import Path
from typing import Any

from .docx_extract import iter_hyperlinks, visible_paragraph_text
from .utilities import IntegrityError, RenderError, normalized_text, sha256_file

try:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.opc.constants import RELATIONSHIP_TYPE
    from docx.shared import Inches, Pt, RGBColor
except ImportError as exc:  # pragma: no cover - guarded by project dependencies
    from .utilities import DependencyError

    raise DependencyError(
        "Python package 'python-docx' is required. Install project dependencies "
        "as documented in README.md."
    ) from exc


HEADLESS_SECTION_HEADINGS = (
    "SUMMARY",
    "EDUCATION & CERTIFICATES",
    "TECHNICAL SKILLS",
    "WORK HISTORY & PROJECTS",
)

_FONT = "Arial"
_BODY_SIZE = 10.5
_CONTACT_SIZE = 9.0
_NAME_SIZE = 14.0
_TEXT_WIDTH_INCHES = 7.0


def _set_run_font(
    run: Any,
    *,
    size: float = _BODY_SIZE,
    bold: bool | None = None,
    italic: bool | None = None,
    color: RGBColor | None = None,
    underline: bool | None = None,
) -> None:
    run.font.name = _FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), _FONT)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    if underline is not None:
        run.underline = underline


def _configure_paragraph(
    paragraph: Any,
    *,
    before: float = 0,
    after: float = 1.0,
    line_spacing: float = 1.08,
) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line_spacing
    fmt.keep_together = True
    fmt.widow_control = True


def _add_bottom_rule(paragraph: Any) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    borders = ppr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        ppr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    borders.append(bottom)


def _add_section_heading(document: Any, text: str) -> None:
    paragraph = document.add_paragraph()
    _configure_paragraph(paragraph, before=3.5, after=1.0, line_spacing=1.0)
    paragraph.paragraph_format.keep_with_next = True
    _add_bottom_rule(paragraph)
    _set_run_font(paragraph.add_run(text), size=_BODY_SIZE, bold=True)


def _add_labelled_paragraph(document: Any, value: dict[str, str]) -> None:
    paragraph = document.add_paragraph()
    _configure_paragraph(paragraph)
    _set_run_font(paragraph.add_run(f"{value['label']}: "), bold=True)
    _set_run_font(paragraph.add_run(value["text"]))


def _add_bullet(document: Any, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    _configure_paragraph(paragraph, after=0.5, line_spacing=1.04)
    paragraph.paragraph_format.left_indent = Inches(0.22)
    paragraph.paragraph_format.first_line_indent = Inches(-0.16)
    paragraph.paragraph_format.keep_with_next = False
    _set_run_font(paragraph.add_run(text))


def _add_entry_heading(
    document: Any,
    *,
    title: str,
    detail: str,
    dates: str | None = None,
) -> None:
    paragraph = document.add_paragraph()
    _configure_paragraph(paragraph, before=1.5, after=0.5, line_spacing=1.0)
    paragraph.paragraph_format.keep_with_next = True
    if dates:
        paragraph.paragraph_format.tab_stops.add_tab_stop(
            Inches(_TEXT_WIDTH_INCHES), WD_TAB_ALIGNMENT.RIGHT
        )
    _set_run_font(paragraph.add_run(title), bold=True, italic=True)
    _set_run_font(paragraph.add_run(f" | {detail}"), italic=True)
    if dates:
        _set_run_font(paragraph.add_run(f"\t{dates}"), italic=True)


def _add_hyperlink(
    paragraph: Any,
    *,
    text: str,
    target: str,
) -> None:
    relationship_id = paragraph.part.relate_to(
        target,
        RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run_element = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")

    fonts = OxmlElement("w:rFonts")
    for attribute in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{attribute}"), _FONT)
    properties.append(fonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "315F7D")
    properties.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(underline)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), str(int(_CONTACT_SIZE * 2)))
    properties.append(size)
    size_complex = OxmlElement("w:szCs")
    size_complex.set(qn("w:val"), str(int(_CONTACT_SIZE * 2)))
    properties.append(size_complex)
    run_element.append(properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run_element.append(text_element)
    hyperlink.append(run_element)
    paragraph._p.append(hyperlink)


def _header_text(extracted_resume: dict[str, Any], content_id: str) -> str:
    for paragraph in extracted_resume.get("paragraphs", []):
        if paragraph.get("content_id") == content_id:
            text = paragraph.get("text")
            if isinstance(text, str) and text.strip():
                return text.strip()
    raise RenderError(f"Extracted master résumé is missing {content_id!r}.")


def _add_contact_header(
    document: Any,
    *,
    contact_text: str,
    links: list[dict[str, Any]],
) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _configure_paragraph(paragraph, after=2.0, line_spacing=1.0)
    cursor = 0
    for link in links:
        display = link.get("text")
        target = link.get("target")
        if not isinstance(display, str) or not display:
            raise RenderError("The master résumé contains an empty contact hyperlink.")
        if not isinstance(target, str) or not target:
            raise RenderError("The master résumé contains a broken contact hyperlink.")
        index = contact_text.find(display, cursor)
        if index < cursor:
            raise RenderError(
                "A contact hyperlink cannot be mapped back to the master résumé header."
            )
        if index > cursor:
            _set_run_font(
                paragraph.add_run(contact_text[cursor:index]),
                size=_CONTACT_SIZE,
            )
        _add_hyperlink(paragraph, text=display, target=target)
        cursor = index + len(display)
    if cursor < len(contact_text):
        _set_run_font(paragraph.add_run(contact_text[cursor:]), size=_CONTACT_SIZE)


def _configure_document(document: Any) -> None:
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.header_distance = Inches(0.2)
    section.footer_distance = Inches(0.2)

    normal = document.styles["Normal"]
    normal.font.name = _FONT
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), _FONT)
    normal.font.size = Pt(_BODY_SIZE)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(1)
    normal.paragraph_format.line_spacing = 1.08
    if "List Bullet" not in document.styles:
        document.styles.add_style("List Bullet", WD_STYLE_TYPE.PARAGRAPH)
    bullet = document.styles["List Bullet"]
    bullet.font.name = _FONT
    bullet._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), _FONT)
    bullet.font.size = Pt(_BODY_SIZE)


def _all_content_strings(tailored_content: dict[str, Any]) -> list[str]:
    education = tailored_content["education"]
    result = [
        tailored_content["professional_summary"],
        education["institution"],
        education["degree_details"],
        education["coursework"]["label"],
        education["coursework"]["text"],
        education["certifications"]["label"],
        education["certifications"]["text"],
    ]
    for group in tailored_content["skill_groups"]:
        result.extend((group["label"], group["text"]))
    for project in tailored_content["projects"]:
        result.extend((project["name"], project["technologies"], *project["bullets"]))
    open_source = tailored_content["open_source"]
    result.extend(
        (open_source["name"], open_source["technologies"], open_source["bullet"])
    )
    experience = tailored_content["experience"]
    result.extend(
        (
            experience["role"],
            experience["employer_location"],
            experience["dates"],
            *experience["bullets"],
        )
    )
    return result


def _expected_paragraph_texts(
    *,
    name: str,
    contact: str,
    tailored_content: dict[str, Any],
) -> list[str]:
    education = tailored_content["education"]
    expected = [
        name,
        contact,
        HEADLESS_SECTION_HEADINGS[0],
        tailored_content["professional_summary"],
        HEADLESS_SECTION_HEADINGS[1],
        f"{education['institution']} | {education['degree_details']}",
        f"{education['coursework']['label']}: {education['coursework']['text']}",
        (
            f"{education['certifications']['label']}: "
            f"{education['certifications']['text']}"
        ),
        HEADLESS_SECTION_HEADINGS[2],
    ]
    expected.extend(
        f"{group['label']}: {group['text']}"
        for group in tailored_content["skill_groups"]
    )
    expected.append(HEADLESS_SECTION_HEADINGS[3])
    experience = tailored_content["experience"]
    expected.append(
        f"{experience['role']} | {experience['employer_location']}\t"
        f"{experience['dates']}"
    )
    expected.extend(experience["bullets"])
    for project in tailored_content["projects"]:
        expected.append(f"{project['name']} | {project['technologies']}")
        expected.extend(project["bullets"])
    open_source = tailored_content["open_source"]
    expected.extend(
        (
            f"{open_source['name']} | {open_source['technologies']}",
            open_source["bullet"],
        )
    )
    return expected


def _validate_rendered_document(
    path: Path,
    *,
    name: str,
    contact: str,
    tailored_content: dict[str, Any],
    expected_links: list[tuple[str, str]],
) -> None:
    document = Document(str(path))
    if len(document.sections) != 1:
        raise RenderError("Headless output must contain exactly one document section.")
    if document.tables:
        raise RenderError("Headless output must remain a single-column, table-free document.")
    actual_paragraphs = [
        visible_paragraph_text(paragraph).strip()
        for paragraph in document.paragraphs
    ]
    expected_paragraphs = _expected_paragraph_texts(
        name=name,
        contact=contact,
        tailored_content=tailored_content,
    )
    if actual_paragraphs != expected_paragraphs:
        raise RenderError(
            "Headless output paragraph content or ordering changed during save."
        )
    visible = "\n".join(visible_paragraph_text(p) for p in document.paragraphs)
    normalized = normalized_text(visible)
    for required in (name, contact, *HEADLESS_SECTION_HEADINGS, *_all_content_strings(tailored_content)):
        if normalized_text(required) not in normalized:
            raise RenderError(f"Headless output is missing required content: {required!r}.")
    actual_links = [
        (str(link["text"]), str(link["target"])) for link in iter_hyperlinks(document)
    ]
    if actual_links != expected_links:
        raise RenderError("Headless output changed a contact hyperlink or target.")
    for paragraph in document.paragraphs:
        for run_element in paragraph._p.xpath(".//w:r"):
            properties = run_element.find(qn("w:rPr"))
            fonts = properties.find(qn("w:rFonts")) if properties is not None else None
            if fonts is None or fonts.get(qn("w:ascii")) != _FONT:
                raise RenderError("Headless output contains a non-Arial text run.")


def render_headless_docx(
    *,
    source_path: Path,
    destination_path: Path,
    tailored_content: dict[str, Any],
    extracted_resume: dict[str, Any],
    expected_source_hash: str,
) -> None:
    """Render grounded content into a deterministic Headless-style résumé.

    Qwen controls wording only. Page geometry, typography, ordering, hyperlinks,
    and structural validation remain deterministic Python responsibilities.
    """
    if source_path.resolve() == destination_path.resolve():
        raise IntegrityError("Refusing to render over the master resume.")
    if destination_path.exists():
        raise IntegrityError(f"Refusing to overwrite existing output: {destination_path}")
    if sha256_file(source_path) != expected_source_hash:
        raise IntegrityError("The master resume changed before rendering began.")

    name = _header_text(extracted_resume, "header.name")
    contact = _header_text(extracted_resume, "header.contact")
    raw_links = extracted_resume.get("document", {}).get("hyperlinks", [])
    links = [link for link in raw_links if link.get("paragraph_index") == 1]
    expected_links = [(str(link["text"]), str(link["target"])) for link in links]

    document = Document()
    _configure_document(document)
    name_paragraph = document.add_paragraph()
    name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _configure_paragraph(name_paragraph, after=0.5, line_spacing=1.0)
    _set_run_font(
        name_paragraph.add_run(name),
        size=_NAME_SIZE,
        bold=True,
    )
    _add_contact_header(document, contact_text=contact, links=links)

    _add_section_heading(document, HEADLESS_SECTION_HEADINGS[0])
    summary = document.add_paragraph()
    _configure_paragraph(summary)
    _set_run_font(summary.add_run(tailored_content["professional_summary"]))

    _add_section_heading(document, HEADLESS_SECTION_HEADINGS[1])
    education = tailored_content["education"]
    degree = document.add_paragraph()
    _configure_paragraph(degree)
    _set_run_font(degree.add_run(education["institution"]), bold=True)
    _set_run_font(degree.add_run(f" | {education['degree_details']}"), italic=True)
    _add_labelled_paragraph(document, education["coursework"])
    _add_labelled_paragraph(document, education["certifications"])

    _add_section_heading(document, HEADLESS_SECTION_HEADINGS[2])
    for group in tailored_content["skill_groups"]:
        _add_labelled_paragraph(document, group)

    _add_section_heading(document, HEADLESS_SECTION_HEADINGS[3])
    experience = tailored_content["experience"]
    _add_entry_heading(
        document,
        title=experience["role"],
        detail=experience["employer_location"],
        dates=experience["dates"],
    )
    for bullet in experience["bullets"]:
        _add_bullet(document, bullet)
    for project in tailored_content["projects"]:
        _add_entry_heading(
            document,
            title=project["name"],
            detail=project["technologies"],
        )
        for bullet in project["bullets"]:
            _add_bullet(document, bullet)
    open_source = tailored_content["open_source"]
    _add_entry_heading(
        document,
        title=open_source["name"],
        detail=open_source["technologies"],
    )
    _add_bullet(document, open_source["bullet"])

    destination_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(destination_path))
    if sha256_file(source_path) != expected_source_hash:
        raise IntegrityError("The master resume changed while the copy was rendered.")
    _validate_rendered_document(
        destination_path,
        name=name,
        contact=contact,
        tailored_content=tailored_content,
        expected_links=expected_links,
    )
