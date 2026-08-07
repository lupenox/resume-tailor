from __future__ import annotations

import os
import shutil
import xml.etree.ElementTree as ET
import zipfile
from pathlib import Path
from typing import Any

from lxml import etree

from resume_tailor.backend.documents.docx_extract import (
    TemplateMap,
    iter_hyperlinks,
    logical_content,
    validate_template,
    visible_paragraph_text,
)
from resume_tailor.backend.utils.utilities import (
    IntegrityError,
    RenderError,
    concise_process_error,
    normalized_text,
    require_executable,
    run_command,
    sha256_file,
)

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph
except ImportError as exc:  # pragma: no cover
    from resume_tailor.backend.utils.utilities import DependencyError

    raise DependencyError(
        "Python package 'python-docx' is required. Install project dependencies "
        "as documented in README.md."
    ) from exc


def _canonical_xml(element: Any | None) -> bytes:
    if element is None:
        return b""
    return etree.tostring(element, method="c14n")


def _formatting_snapshot(document: Any) -> list[dict[str, Any]]:
    snapshot: list[dict[str, Any]] = []
    for paragraph in document.paragraphs:
        runs: list[tuple[bool, bytes]] = []
        for run in paragraph._p.xpath(".//w:r"):
            runs.append(
                (
                    run.getparent().tag == qn("w:hyperlink"),
                    _canonical_xml(run.find(qn("w:rPr"))),
                )
            )
        snapshot.append(
            {
                "paragraph_properties": _canonical_xml(paragraph._p.pPr),
                "run_properties": runs,
            }
        )
    return snapshot


def _geometry_snapshot(document: Any) -> tuple[int, ...]:
    section = document.sections[0]
    values = (
        section.page_width,
        section.page_height,
        section.top_margin,
        section.right_margin,
        section.bottom_margin,
        section.left_margin,
        section.header_distance,
        section.footer_distance,
    )
    return tuple(int(value) for value in values)


def _hyperlink_snapshot(document: Any) -> list[tuple[int, str, str]]:
    return [
        (
            int(link["paragraph_index"]),
            str(link["text"]),
            str(link["target"]),
        )
        for link in iter_hyperlinks(document)
    ]


_PRESERVED_PACKAGE_PARTS = (
    "word/_rels/document.xml.rels",
    "word/styles.xml",
    "word/stylesWithEffects.xml",
    "word/numbering.xml",
    "word/settings.xml",
    "word/webSettings.xml",
    "word/fontTable.xml",
    "word/theme/theme1.xml",
)


def _package_part_snapshot(path: Path) -> dict[str, bytes]:
    try:
        with zipfile.ZipFile(path) as archive:
            return {
                name: archive.read(name)
                for name in _PRESERVED_PACKAGE_PARTS
                if name in archive.namelist()
            }
    except (OSError, zipfile.BadZipFile, KeyError) as exc:
        raise RenderError(f"Could not inspect preserved DOCX package parts: {exc}") from exc


def _replace_plain_text(paragraph: Paragraph, value: str, *, label: str) -> None:
    runs = paragraph.runs
    if not runs:
        raise RenderError(f"Cannot replace {label}: paragraph has no direct text run.")
    runs[0].text = value
    for run in runs[1:]:
        run.text = ""


def _replace_labelled(
    paragraph: Paragraph,
    value: dict[str, str],
    *,
    label: str,
) -> None:
    runs = paragraph.runs
    if len(runs) < 2:
        raise RenderError(
            f"Cannot replace {label}: expected separate label and body runs."
        )
    if runs[0].bold is not True:
        raise RenderError(f"Cannot replace {label}: the label run is no longer bold.")
    runs[0].text = f"{value['label']}: "
    runs[1].text = value["text"]
    for run in runs[2:]:
        run.text = ""


def _replace_project_heading(
    paragraph: Paragraph,
    project: dict[str, Any],
    *,
    label: str,
) -> None:
    runs = paragraph.runs
    if len(runs) < 3 or runs[0].bold is not True or runs[-1].italic is not True:
        raise RenderError(
            f"Cannot replace {label}: expected bold name, separator, and italic "
            "technology runs."
        )
    runs[0].text = project["name"]
    runs[-1].text = project["technologies"]


def render_tailored_docx(
    *,
    source_path: Path,
    destination_path: Path,
    tailored_content: dict[str, Any],
    expected_source_hash: str,
) -> None:
    source_resolved = source_path.resolve()
    destination_resolved = destination_path.resolve()
    if source_resolved == destination_resolved:
        raise IntegrityError("Refusing to render over the master resume.")
    if destination_path.exists():
        raise IntegrityError(f"Refusing to overwrite existing output: {destination_path}")
    if sha256_file(source_path) != expected_source_hash:
        raise IntegrityError("The master resume changed before rendering began.")

    source_document = Document(str(source_path))
    source_map = validate_template(source_document)
    source_formatting = _formatting_snapshot(source_document)
    source_geometry = _geometry_snapshot(source_document)
    source_hyperlinks = _hyperlink_snapshot(source_document)
    source_package_parts = _package_part_snapshot(source_path)
    source_contact = [
        visible_paragraph_text(source_document.paragraphs[index])
        for index in (source_map.name, source_map.contact)
    ]

    destination_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(source_path, destination_path)
    working_document = Document(str(destination_path))
    mapping = validate_template(working_document)
    paragraphs = working_document.paragraphs

    _replace_plain_text(
        paragraphs[mapping.summary],
        tailored_content["professional_summary"],
        label="professional summary",
    )
    _replace_labelled(
        paragraphs[mapping.education_coursework],
        tailored_content["education"]["coursework"],
        label="education coursework",
    )
    _replace_labelled(
        paragraphs[mapping.education_certifications],
        tailored_content["education"]["certifications"],
        label="education certifications",
    )
    for index, paragraph_index in enumerate(mapping.skill_groups):
        _replace_labelled(
            paragraphs[paragraph_index],
            tailored_content["skill_groups"][index],
            label=f"skill group {index + 1}",
        )
    for project_index, project_map in enumerate(mapping.projects):
        project_content = tailored_content["projects"][project_index]
        _replace_project_heading(
            paragraphs[project_map.heading],
            project_content,
            label=f"project {project_index + 1}",
        )
        for bullet_index, paragraph_index in enumerate(project_map.bullets):
            _replace_plain_text(
                paragraphs[paragraph_index],
                project_content["bullets"][bullet_index],
                label=f"project {project_index + 1} bullet {bullet_index + 1}",
            )
    _replace_plain_text(
        paragraphs[mapping.open_source_bullet],
        tailored_content["open_source"]["bullet"],
        label="open-source bullet",
    )
    for bullet_index, paragraph_index in enumerate(mapping.experience_bullets):
        _replace_plain_text(
            paragraphs[paragraph_index],
            tailored_content["experience"]["bullets"][bullet_index],
            label=f"experience bullet {bullet_index + 1}",
        )

    working_document.save(str(destination_path))
    if sha256_file(source_path) != expected_source_hash:
        raise IntegrityError("The master resume changed while the copy was rendered.")

    rendered_document = Document(str(destination_path))
    rendered_map = validate_template(rendered_document)
    if rendered_map != source_map:
        raise RenderError("The rendered document's semantic structure changed.")
    if _geometry_snapshot(rendered_document) != source_geometry:
        raise RenderError("The rendered document's page geometry or margins changed.")
    if _formatting_snapshot(rendered_document) != source_formatting:
        raise RenderError(
            "Run or paragraph formatting changed outside text nodes; output was preserved "
            "for inspection."
        )
    if _hyperlink_snapshot(rendered_document) != source_hyperlinks:
        raise RenderError("A hyperlink or hyperlink target changed during rendering.")
    if _package_part_snapshot(destination_path) != source_package_parts:
        raise RenderError(
            "A preserved styles, numbering, settings, font, theme, or relationship "
            "package part changed during rendering."
        )
    rendered_contact = [
        visible_paragraph_text(rendered_document.paragraphs[index])
        for index in (rendered_map.name, rendered_map.contact)
    ]
    if rendered_contact != source_contact:
        raise RenderError("Header or contact information changed during rendering.")
    if logical_content(rendered_document, rendered_map) != tailored_content:
        raise RenderError(
            "Post-save DOCX content does not exactly match the approved structured content."
        )


def _parse_pdf_pages(pdfinfo_output: str) -> int:
    for line in pdfinfo_output.splitlines():
        if line.startswith("Pages:"):
            try:
                return int(line.split(":", 1)[1].strip())
            except ValueError as exc:
                raise RenderError("pdfinfo returned an invalid page count.") from exc
    raise RenderError("pdfinfo did not report a page count.")


def _validate_bbox(path: Path) -> None:
    try:
        root = ET.parse(path).getroot()
    except (OSError, ET.ParseError) as exc:
        raise RenderError(f"Could not parse Poppler bounding-box output: {exc}") from exc
    pages = [node for node in root.iter() if node.tag.rsplit("}", 1)[-1] == "page"]
    if len(pages) != 1:
        raise RenderError(f"Bounding-box output contains {len(pages)} pages, expected one.")
    page = pages[0]
    width = float(page.attrib["width"])
    height = float(page.attrib["height"])
    lines: list[tuple[float, float, float, float]] = []
    for node in page.iter():
        local_name = node.tag.rsplit("}", 1)[-1]
        if local_name not in {"word", "line"}:
            continue
        try:
            box = tuple(
                float(node.attrib[key]) for key in ("xMin", "yMin", "xMax", "yMax")
            )
        except (KeyError, ValueError) as exc:
            raise RenderError("Poppler returned an invalid text bounding box.") from exc
        x_min, y_min, x_max, y_max = box
        if x_min < -0.5 or y_min < -0.5 or x_max > width + 0.5 or y_max > height + 0.5:
            raise RenderError(
                "PDF text extends outside the page boundary, indicating clipping."
            )
        if local_name == "line":
            lines.append(box)

    for index, first in enumerate(lines):
        for second in lines[index + 1 :]:
            horizontal_overlap = min(first[2], second[2]) - max(first[0], second[0])
            vertical_overlap = min(first[3], second[3]) - max(first[1], second[1])
            smaller_height = min(first[3] - first[1], second[3] - second[1])
            if horizontal_overlap > 2 and vertical_overlap > max(1.5, smaller_height * 0.55):
                if abs(first[1] - second[1]) > 0.5:
                    raise RenderError(
                        "Poppler detected overlapping text lines in the exported PDF."
                    )


def export_and_validate_pdf(
    *,
    docx_path: Path,
    pdf_path: Path,
    preview_path: Path,
    working_directory: Path,
    required_text: list[str],
    timeout_seconds: int = 120,
) -> str:
    libreoffice = require_executable("libreoffice")
    pdfinfo = require_executable("pdfinfo")
    pdftotext = require_executable("pdftotext")
    pdftoppm = require_executable("pdftoppm")
    if pdf_path.exists() or preview_path.exists():
        raise IntegrityError("Refusing to overwrite an existing PDF or preview.")

    profile = working_directory / "libreoffice-profile"
    temporary = working_directory / "tmp"
    export_directory = working_directory / "pdf-export"
    bbox_path = working_directory / "pdf-bbox.html"
    for directory in (profile, temporary, export_directory):
        directory.mkdir(parents=True, exist_ok=True)
    environment = os.environ.copy()
    environment["TMPDIR"] = str(temporary)
    environment["SAL_USE_VCLPLUGIN"] = "svp"
    conversion = run_command(
        [
            libreoffice,
            f"-env:UserInstallation={profile.resolve().as_uri()}",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(export_directory),
            str(docx_path.resolve()),
        ],
        cwd=working_directory,
        timeout_seconds=timeout_seconds,
        env=environment,
    )
    if conversion.returncode != 0:
        raise RenderError(concise_process_error(conversion, "LibreOffice export"))
    generated_pdf = export_directory / f"{docx_path.stem}.pdf"
    if not generated_pdf.is_file():
        raise RenderError("LibreOffice did not create the expected PDF.")
    generated_pdf.replace(pdf_path)

    info = run_command(
        [pdfinfo, str(pdf_path)],
        cwd=working_directory,
        timeout_seconds=30,
    )
    if info.returncode != 0:
        raise RenderError(concise_process_error(info, "pdfinfo"))
    pages = _parse_pdf_pages(info.stdout)
    if pages != 1:
        raise RenderError(
            f"Tailored PDF is {pages} pages; exactly one is required. No fonts or "
            "margins were changed. Shorten approved content and rerun."
        )

    text_result = run_command(
        [pdftotext, "-layout", str(pdf_path), "-"],
        cwd=working_directory,
        timeout_seconds=30,
    )
    if text_result.returncode != 0:
        raise RenderError(concise_process_error(text_result, "pdftotext"))
    extracted_text = text_result.stdout
    if not extracted_text.strip():
        raise RenderError("The exported PDF has no extractable text.")
    if "\ufffd" in extracted_text or "\x00" in extracted_text:
        raise RenderError("The exported PDF contains a missing/replacement glyph.")
    normalized_pdf = normalized_text(extracted_text)
    missing = [
        text
        for text in required_text
        if normalized_text(text) not in normalized_pdf
    ]
    if missing:
        raise RenderError(
            "The exported PDF is missing required header/section text: "
            + ", ".join(repr(item) for item in missing)
        )

    bbox = run_command(
        [pdftotext, "-bbox-layout", str(pdf_path), str(bbox_path)],
        cwd=working_directory,
        timeout_seconds=30,
    )
    if bbox.returncode != 0:
        raise RenderError(concise_process_error(bbox, "pdftotext bounding-box check"))
    _validate_bbox(bbox_path)

    preview_base = preview_path.with_suffix("")
    raster = run_command(
        [
            pdftoppm,
            "-f",
            "1",
            "-singlefile",
            "-png",
            "-r",
            "144",
            str(pdf_path),
            str(preview_base),
        ],
        cwd=working_directory,
        timeout_seconds=timeout_seconds,
    )
    if raster.returncode != 0:
        raise RenderError(concise_process_error(raster, "pdftoppm"))
    if not preview_path.is_file() or preview_path.stat().st_size == 0:
        raise RenderError("Poppler did not create a nonempty preview PNG.")
    return extracted_text
