from __future__ import annotations

import re
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Any, Iterable

from .character_budget import calculate_content_budget
from .utilities import DependencyError, TemplateError, sha256_file

try:
    from docx import Document
    from docx.document import Document as DocumentObject
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph
except ImportError as exc:  # pragma: no cover - exercised by dependency error path
    raise DependencyError(
        "Python package 'python-docx' is required. Install project dependencies "
        "as documented in README.md."
    ) from exc


SECTION_HEADINGS = (
    "OBJECTIVE / SUMMARY",
    "EDUCATION & CERTIFICATIONS",
    "TECHNICAL SKILLS",
    "AI ENGINEERING PROJECTS",
    "OPEN SOURCE CONTRIBUTION",
    "EXPERIENCE",
)
EXPECTED_PROJECT_BULLETS = (3, 4, 3)

_EDITABLE_SOURCE_IDS = {
    "professional_summary",
    "education.coursework",
    "education.certifications",
    "open_source.bullet",
}
_EDITABLE_SOURCE_PREFIXES = (
    "skill_groups.",
    "projects.",
    "experience.bullets.",
)


@dataclass(frozen=True)
class ProjectMap:
    heading: int
    bullets: tuple[int, ...]


@dataclass(frozen=True)
class TemplateMap:
    name: int
    contact: int
    section_headings: dict[str, int]
    summary: int
    education_degree: int
    education_coursework: int
    education_certifications: int
    skill_groups: tuple[int, int, int]
    projects: tuple[ProjectMap, ProjectMap, ProjectMap]
    open_source_heading: int
    open_source_bullet: int
    experience_heading: int
    experience_bullets: tuple[int, ...]

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)


def visible_paragraph_text(paragraph: Paragraph) -> str:
    """Return text including runs nested inside hyperlink elements."""
    pieces: list[str] = []
    for element in paragraph._p.iter():
        if element.tag == qn("w:t") and element.text:
            pieces.append(element.text)
        elif element.tag == qn("w:tab"):
            pieces.append("\t")
        elif element.tag in {qn("w:br"), qn("w:cr")}:
            pieces.append("\n")
    return "".join(pieces)


def iter_hyperlinks(document: DocumentObject) -> Iterable[dict[str, Any]]:
    for index, paragraph in enumerate(document.paragraphs):
        for hyperlink in paragraph._p.xpath(".//w:hyperlink"):
            relationship_id = hyperlink.get(qn("r:id"))
            text = "".join(
                node.text or "" for node in hyperlink.iter() if node.tag == qn("w:t")
            )
            target = None
            if relationship_id:
                relationship = document.part.rels.get(relationship_id)
                target = relationship.target_ref if relationship is not None else None
            yield {
                "paragraph_index": index,
                "relationship_id": relationship_id,
                "text": text,
                "target": target,
            }


def _is_list(paragraph: Paragraph) -> bool:
    style_name = paragraph.style.name if paragraph.style is not None else ""
    return style_name.startswith("List")


def _has_bottom_rule(paragraph: Paragraph) -> bool:
    return bool(paragraph._p.xpath("./w:pPr/w:pBdr/w:bottom"))


def _direct_run_descriptors(paragraph: Paragraph) -> list[dict[str, Any]]:
    descriptors: list[dict[str, Any]] = []
    for run_element in paragraph._p.xpath(".//w:r"):
        properties = run_element.find(qn("w:rPr"))
        text = "".join(
            node.text or "" for node in run_element.iter() if node.tag == qn("w:t")
        )

        def has(tag: str) -> bool:
            return properties is not None and properties.find(qn(tag)) is not None

        size = None
        color = None
        font = None
        if properties is not None:
            size_node = properties.find(qn("w:sz"))
            color_node = properties.find(qn("w:color"))
            fonts_node = properties.find(qn("w:rFonts"))
            if size_node is not None and size_node.get(qn("w:val")):
                size = int(size_node.get(qn("w:val"))) / 2
            if color_node is not None:
                color = color_node.get(qn("w:val"))
            if fonts_node is not None:
                font = fonts_node.get(qn("w:ascii"))
        descriptors.append(
            {
                "text": text,
                "bold": has("w:b"),
                "italic": has("w:i"),
                "underline": has("w:u"),
                "size_points": size,
                "font": font,
                "color": color,
                "inside_hyperlink": run_element.getparent().tag == qn("w:hyperlink"),
            }
        )
    return descriptors


def _spacing_descriptor(paragraph: Paragraph) -> dict[str, Any]:
    spacing = paragraph._p.find(qn("w:pPr"))
    if spacing is None:
        return {}
    node = spacing.find(qn("w:spacing"))
    if node is None:
        return {}
    return {
        key: node.get(qn(f"w:{key}"))
        for key in ("before", "after", "line", "lineRule")
        if node.get(qn(f"w:{key}")) is not None
    }


def _content_budget(text: str) -> dict[str, int]:
    return calculate_content_budget(text)


def _section_map(document: DocumentObject) -> dict[str, int]:
    texts = [visible_paragraph_text(paragraph).strip() for paragraph in document.paragraphs]
    result: dict[str, int] = {}
    for heading in SECTION_HEADINGS:
        occurrences = [index for index, text in enumerate(texts) if text == heading]
        if len(occurrences) != 1:
            raise TemplateError(
                f"Template drift: expected one '{heading}' heading, found "
                f"{len(occurrences)}."
            )
        result[heading] = occurrences[0]
    indexes = [result[heading] for heading in SECTION_HEADINGS]
    if indexes != sorted(indexes):
        raise TemplateError("Template drift: major resume sections are out of order.")
    return result


def _span(section_map: dict[str, int], current: str, following: str | None, total: int) -> range:
    start = section_map[current] + 1
    end = section_map[following] if following else total
    return range(start, end)


def build_template_map(document: DocumentObject) -> TemplateMap:
    paragraphs = document.paragraphs
    section_map = _section_map(document)

    summary_span = list(
        _span(section_map, SECTION_HEADINGS[0], SECTION_HEADINGS[1], len(paragraphs))
    )
    education_span = list(
        _span(section_map, SECTION_HEADINGS[1], SECTION_HEADINGS[2], len(paragraphs))
    )
    skills_span = list(
        _span(section_map, SECTION_HEADINGS[2], SECTION_HEADINGS[3], len(paragraphs))
    )
    projects_span = list(
        _span(section_map, SECTION_HEADINGS[3], SECTION_HEADINGS[4], len(paragraphs))
    )
    open_source_span = list(
        _span(section_map, SECTION_HEADINGS[4], SECTION_HEADINGS[5], len(paragraphs))
    )
    experience_span = list(
        _span(section_map, SECTION_HEADINGS[5], None, len(paragraphs))
    )

    if len(summary_span) != 1:
        raise TemplateError("Template drift: summary must contain exactly one paragraph.")
    if len(education_span) != 3 or any(
        not _is_list(paragraphs[index]) for index in education_span[1:]
    ):
        raise TemplateError(
            "Template drift: education must contain one degree line and two bullets."
        )
    if len(skills_span) != 3 or any(
        not _is_list(paragraphs[index]) for index in skills_span
    ):
        raise TemplateError("Template drift: expected exactly three technical-skill bullets.")

    project_maps: list[ProjectMap] = []
    cursor = 0
    while cursor < len(projects_span):
        heading = projects_span[cursor]
        if _is_list(paragraphs[heading]):
            raise TemplateError("Template drift: project heading was replaced by a list paragraph.")
        cursor += 1
        bullets: list[int] = []
        while cursor < len(projects_span) and _is_list(paragraphs[projects_span[cursor]]):
            bullets.append(projects_span[cursor])
            cursor += 1
        if not bullets:
            raise TemplateError("Template drift: every project must retain at least one bullet.")
        project_maps.append(ProjectMap(heading, tuple(bullets)))
    counts = tuple(len(project.bullets) for project in project_maps)
    if len(project_maps) != 3 or counts != EXPECTED_PROJECT_BULLETS:
        raise TemplateError(
            "Template drift: expected three projects with bullet counts "
            f"{EXPECTED_PROJECT_BULLETS}, found {counts}."
        )

    if (
        len(open_source_span) != 2
        or _is_list(paragraphs[open_source_span[0]])
        or not _is_list(paragraphs[open_source_span[1]])
    ):
        raise TemplateError(
            "Template drift: open-source section must contain one heading and one bullet."
        )
    if (
        len(experience_span) != 2
        or _is_list(paragraphs[experience_span[0]])
        or not _is_list(paragraphs[experience_span[1]])
    ):
        raise TemplateError(
            "Template drift: experience must contain one position heading and one bullet."
        )

    return TemplateMap(
        name=0,
        contact=1,
        section_headings=section_map,
        summary=summary_span[0],
        education_degree=education_span[0],
        education_coursework=education_span[1],
        education_certifications=education_span[2],
        skill_groups=tuple(skills_span),  # type: ignore[arg-type]
        projects=tuple(project_maps),  # type: ignore[arg-type]
        open_source_heading=open_source_span[0],
        open_source_bullet=open_source_span[1],
        experience_heading=experience_span[0],
        experience_bullets=(experience_span[1],),
    )


def _expect_run_pattern(
    paragraph: Paragraph,
    *,
    label: str,
    first_bold: bool = False,
    last_italic: bool = False,
) -> None:
    descriptors = _direct_run_descriptors(paragraph)
    if not descriptors:
        raise TemplateError(f"Template drift: {label} has no text runs.")
    if first_bold and not descriptors[0]["bold"]:
        raise TemplateError(f"Template drift: {label} no longer begins with a bold run.")
    if last_italic and not descriptors[-1]["italic"]:
        raise TemplateError(f"Template drift: {label} no longer ends with an italic run.")


def validate_template(document: DocumentObject) -> TemplateMap:
    if len(document.sections) != 1:
        raise TemplateError(
            f"Template drift: expected one section, found {len(document.sections)}."
        )
    if len(document.paragraphs) != 32:
        raise TemplateError(
            f"Template drift: expected 32 paragraphs, found {len(document.paragraphs)}."
        )
    if document.tables:
        raise TemplateError(
            f"Template drift: expected zero tables, found {len(document.tables)}."
        )

    section = document.sections[0]
    geometry = {
        "page_width": section.page_width.inches,
        "page_height": section.page_height.inches,
        "top": section.top_margin.inches,
        "right": section.right_margin.inches,
        "bottom": section.bottom_margin.inches,
        "left": section.left_margin.inches,
    }
    expected = {
        "page_width": 8.5,
        "page_height": 11.0,
        "top": 691 / 1440,
        "right": 0.5,
        "bottom": 691 / 1440,
        "left": 0.5,
    }
    for key, expected_value in expected.items():
        if abs(geometry[key] - expected_value) > 0.03:
            raise TemplateError(
                f"Template drift: {key.replace('_', ' ')} is {geometry[key]:.3f}in; "
                f"expected approximately {expected_value:.3f}in."
            )

    mapping = build_template_map(document)
    paragraphs = document.paragraphs
    for heading, index in mapping.section_headings.items():
        if not _has_bottom_rule(paragraphs[index]):
            raise TemplateError(
                f"Template drift: section heading '{heading}' lost its divider rule."
            )
        _expect_run_pattern(paragraphs[index], label=heading, first_bold=True)

    for label, index in (
        ("education coursework", mapping.education_coursework),
        ("education certifications", mapping.education_certifications),
        *((f"skill group {number + 1}", index) for number, index in enumerate(mapping.skill_groups)),
    ):
        _expect_run_pattern(paragraphs[index], label=label, first_bold=True)

    _expect_run_pattern(
        paragraphs[mapping.education_degree],
        label="education degree",
        first_bold=True,
    )
    for number, project in enumerate(mapping.projects, start=1):
        _expect_run_pattern(
            paragraphs[project.heading],
            label=f"project heading {number}",
            first_bold=True,
            last_italic=True,
        )
    _expect_run_pattern(
        paragraphs[mapping.open_source_heading],
        label="open-source heading",
        first_bold=True,
        last_italic=True,
    )
    _expect_run_pattern(
        paragraphs[mapping.experience_heading],
        label="experience heading",
        first_bold=True,
        last_italic=True,
    )

    hyperlinks = list(iter_hyperlinks(document))
    if len(hyperlinks) != 6:
        raise TemplateError(
            f"Template drift: expected six hyperlinks, found {len(hyperlinks)}."
        )
    if any(not link["relationship_id"] or not link["target"] for link in hyperlinks):
        raise TemplateError("Template drift: a hyperlink relationship is broken.")
    if not any(
        "github.com/" in str(link["target"])
        and "/pull/" in str(link["target"])
        for link in hyperlinks
    ):
        raise TemplateError("Template drift: the open-source pull-request hyperlink is missing.")

    for paragraph in paragraphs:
        for run in _direct_run_descriptors(paragraph):
            size = run["size_points"]
            if size is not None and size < 9:
                raise TemplateError(
                    f"Template contains {size:g}pt text; the minimum permitted size is 9pt."
                )
    return mapping


def _split_labelled(paragraph: Paragraph) -> dict[str, str]:
    text = visible_paragraph_text(paragraph).strip()
    if ":" not in text:
        raise TemplateError(f"Expected a bold label followed by a colon: {text!r}")
    label, body = text.split(":", 1)
    if not label.strip() or not body.strip():
        raise TemplateError(f"Labelled paragraph is incomplete: {text!r}")
    return {"label": label.strip(), "text": body.strip()}


def _project_content(document: DocumentObject, project: ProjectMap) -> dict[str, Any]:
    paragraph = document.paragraphs[project.heading]
    runs = paragraph.runs
    if len(runs) < 3:
        raise TemplateError("Project heading no longer has name, separator, and technology runs.")
    return {
        "name": runs[0].text.strip(),
        "technologies": runs[-1].text.strip(),
        "bullets": [
            visible_paragraph_text(document.paragraphs[index]).strip()
            for index in project.bullets
        ],
    }


def logical_content(document: DocumentObject, mapping: TemplateMap) -> dict[str, Any]:
    paragraphs = document.paragraphs
    degree_runs = paragraphs[mapping.education_degree].runs
    if len(degree_runs) < 2:
        raise TemplateError("Education line no longer has separate institution and degree runs.")
    open_source_text = visible_paragraph_text(
        paragraphs[mapping.open_source_heading]
    ).strip()
    if " | " not in open_source_text:
        raise TemplateError("Open-source heading no longer contains the expected separator.")
    open_source_name, open_source_technologies = open_source_text.split(" | ", 1)

    experience_runs = paragraphs[mapping.experience_heading].runs
    if len(experience_runs) < 3:
        raise TemplateError("Experience heading no longer has role, employer, and date runs.")
    employer = experience_runs[1].text.strip()
    if employer.startswith("|"):
        employer = employer[1:].strip()
    dates = experience_runs[-1].text.strip()
    if dates.startswith("(") and dates.endswith(")"):
        dates = dates[1:-1].strip()

    return {
        "professional_summary": visible_paragraph_text(
            paragraphs[mapping.summary]
        ).strip(),
        "education": {
            "institution": degree_runs[0].text.strip(),
            "degree_details": "".join(run.text for run in degree_runs[1:]).strip(" |"),
            "coursework": _split_labelled(paragraphs[mapping.education_coursework]),
            "certifications": _split_labelled(
                paragraphs[mapping.education_certifications]
            ),
        },
        "skill_groups": [
            _split_labelled(paragraphs[index]) for index in mapping.skill_groups
        ],
        "projects": [
            _project_content(document, project) for project in mapping.projects
        ],
        "open_source": {
            "name": open_source_name.strip(),
            "technologies": open_source_technologies.strip(),
            "bullet": visible_paragraph_text(
                paragraphs[mapping.open_source_bullet]
            ).strip(),
        },
        "experience": {
            "role": experience_runs[0].text.strip(),
            "employer_location": employer,
            "dates": dates,
            "bullets": [
                visible_paragraph_text(paragraphs[index]).strip()
                for index in mapping.experience_bullets
            ],
        },
    }


def _content_ids(mapping: TemplateMap) -> dict[int, str]:
    result = {
        mapping.name: "header.name",
        mapping.contact: "header.contact",
        mapping.summary: "professional_summary",
        mapping.education_degree: "education.degree",
        mapping.education_coursework: "education.coursework",
        mapping.education_certifications: "education.certifications",
        mapping.open_source_heading: "open_source.heading",
        mapping.open_source_bullet: "open_source.bullet",
        mapping.experience_heading: "experience.heading",
    }
    for number, index in enumerate(mapping.skill_groups):
        result[index] = f"skill_groups.{number}"
    for project_number, project in enumerate(mapping.projects):
        result[project.heading] = f"projects.{project_number}.heading"
        for bullet_number, index in enumerate(project.bullets):
            result[index] = f"projects.{project_number}.bullets.{bullet_number}"
    for bullet_number, index in enumerate(mapping.experience_bullets):
        result[index] = f"experience.bullets.{bullet_number}"
    for heading, index in mapping.section_headings.items():
        result[index] = f"section.{re.sub(r'[^a-z]+', '_', heading.casefold()).strip('_')}"
    return result


def _is_editable_source_id(source_id: str) -> bool:
    if source_id in _EDITABLE_SOURCE_IDS:
        return True
    if source_id.startswith("projects."):
        return ".bullets." in source_id
    return source_id.startswith(_EDITABLE_SOURCE_PREFIXES)


def source_blocks_from_paragraphs(
    paragraphs: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    """Build the immutable, deterministic source catalog used by model stages."""
    blocks: list[dict[str, Any]] = []
    seen: set[str] = set()
    section_context = "Header"
    for paragraph in paragraphs:
        source_id = str(paragraph.get("content_id", "")).strip()
        if not source_id:
            raise TemplateError("Extracted paragraph is missing its deterministic source ID.")
        if source_id in seen:
            raise TemplateError(f"Duplicate extracted source ID: {source_id!r}.")
        seen.add(source_id)
        exact_text = paragraph.get("text")
        if not isinstance(exact_text, str):
            raise TemplateError(f"Extracted source {source_id!r} has no exact text.")

        if source_id.startswith("section."):
            section_context = exact_text
            block_kind = "section_heading"
        elif source_id.startswith("header."):
            block_kind = "header"
        elif paragraph.get("is_list") is True:
            block_kind = "list_item"
        else:
            block_kind = "paragraph"

        context = "Header" if block_kind == "header" else section_context
        evidence_allowed = block_kind not in {"header", "section_heading"}
        blocks.append(
            {
                "source_id": source_id,
                "section_context": context,
                "block_kind": block_kind,
                "exact_text": exact_text,
                "evidence_allowed": evidence_allowed,
                "editable": evidence_allowed and _is_editable_source_id(source_id),
            }
        )
    return blocks


def extract_resume(path: Path) -> tuple[dict[str, Any], TemplateMap]:
    try:
        document = Document(str(path))
    except (OSError, ValueError) as exc:
        raise TemplateError(f"Cannot open DOCX template {path}: {exc}") from exc
    mapping = validate_template(document)
    content_ids = _content_ids(mapping)
    section = document.sections[0]
    hyperlinks = list(iter_hyperlinks(document))
    paragraph_records: list[dict[str, Any]] = []
    for index, paragraph in enumerate(document.paragraphs):
        text = visible_paragraph_text(paragraph).strip()
        ppr = paragraph._p.pPr
        num_pr = ppr.numPr if ppr is not None else None
        paragraph_records.append(
            {
                "index": index,
                "content_id": content_ids.get(index, f"paragraph.{index}"),
                "text": text,
                "style": paragraph.style.name if paragraph.style is not None else None,
                "is_list": _is_list(paragraph),
                "numbering": {
                    "number_id": str(num_pr.numId.val),
                    "level": str(num_pr.ilvl.val),
                }
                if num_pr is not None
                else None,
                "has_bottom_rule": _has_bottom_rule(paragraph),
                "spacing": _spacing_descriptor(paragraph),
                "runs": _direct_run_descriptors(paragraph),
                "content_budget": _content_budget(text),
            }
        )

    extracted = {
        "source": {
            "filename": path.name,
            "sha256": sha256_file(path),
        },
        "document": {
            "sections": len(document.sections),
            "paragraphs": len(document.paragraphs),
            "tables": len(document.tables),
            "page": {
                "width_inches": round(section.page_width.inches, 4),
                "height_inches": round(section.page_height.inches, 4),
                "top_margin_inches": round(section.top_margin.inches, 4),
                "right_margin_inches": round(section.right_margin.inches, 4),
                "bottom_margin_inches": round(section.bottom_margin.inches, 4),
                "left_margin_inches": round(section.left_margin.inches, 4),
            },
            "hyperlinks": hyperlinks,
        },
        "template_map": mapping.to_dict(),
        "content": logical_content(document, mapping),
        "paragraphs": paragraph_records,
        "source_blocks": source_blocks_from_paragraphs(paragraph_records),
    }
    return extracted, mapping


def extract_docx_text(path: Path) -> str:
    try:
        document = Document(str(path))
    except (OSError, ValueError) as exc:
        raise TemplateError(f"Cannot open DOCX {path}: {exc}") from exc
    return "\n".join(
        visible_paragraph_text(paragraph).strip()
        for paragraph in document.paragraphs
        if visible_paragraph_text(paragraph).strip()
    )
