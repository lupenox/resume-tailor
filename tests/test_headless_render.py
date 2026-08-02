from __future__ import annotations

from pathlib import Path

from docx import Document

from resume_tailor.docx_extract import extract_resume, iter_hyperlinks
from resume_tailor.headless_render import (
    HEADLESS_SECTION_HEADINGS,
    render_headless_docx,
)
from resume_tailor.utilities import sha256_file


def test_headless_renderer_preserves_grounded_content_header_and_links(
    master_resume: Path,
    tmp_path: Path,
) -> None:
    extracted, _ = extract_resume(master_resume)
    before = sha256_file(master_resume)
    output = tmp_path / "headless.docx"

    render_headless_docx(
        source_path=master_resume,
        destination_path=output,
        tailored_content=extracted["content"],
        extracted_resume=extracted,
        expected_source_hash=before,
    )

    assert output.is_file()
    assert sha256_file(master_resume) == before
    document = Document(output)
    assert len(document.sections) == 1
    assert document.tables == []
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert all(heading in text for heading in HEADLESS_SECTION_HEADINGS)
    assert "SAMPLE CANDIDATE" in text
    expected_links = [
        (str(link["text"]), str(link["target"]))
        for link in extracted["document"]["hyperlinks"]
    ]
    actual_links = [
        (str(link["text"]), str(link["target"]))
        for link in iter_hyperlinks(document)
    ]
    assert actual_links == expected_links
    assert all(
        run.font.name == "Arial"
        for paragraph in document.paragraphs
        for run in paragraph.runs
    )
