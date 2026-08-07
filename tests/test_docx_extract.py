from __future__ import annotations

import shutil
from pathlib import Path

import pytest
from docx import Document

from resume_tailor.backend.documents.docx_extract import (
    EXPECTED_PROJECT_BULLETS,
    extract_resume,
    validate_template,
)
from resume_tailor.backend.utils.utilities import TemplateError


def test_known_master_structure(master_resume: Path) -> None:
    extracted, mapping = extract_resume(master_resume)
    document = extracted["document"]
    assert document["sections"] == 1
    assert document["paragraphs"] == 32
    assert document["tables"] == 0
    assert document["page"]["width_inches"] == 8.5
    assert document["page"]["height_inches"] == 11.0
    assert document["page"]["top_margin_inches"] == pytest.approx(0.4799, abs=0.001)
    assert document["page"]["left_margin_inches"] == 0.5
    assert len(document["hyperlinks"]) == 6
    assert tuple(len(project.bullets) for project in mapping.projects) == EXPECTED_PROJECT_BULLETS
    assert len(extracted["source_blocks"]) == 32
    assert len({block["source_id"] for block in extracted["source_blocks"]}) == 32
    summary = next(
        block
        for block in extracted["source_blocks"]
        if block["source_id"] == "professional_summary"
    )
    assert summary["editable"] is True
    assert summary["evidence_allowed"] is True
    section = next(
        block
        for block in extracted["source_blocks"]
        if block["source_id"] == "section.objective_summary"
    )
    assert section["evidence_allowed"] is False
    assert section["editable"] is False


def test_synthetic_list_bullet_and_run_formatting(master_resume: Path) -> None:
    extracted, _ = extract_resume(master_resume)
    list_paragraphs = [
        paragraph for paragraph in extracted["paragraphs"] if paragraph["is_list"]
    ]
    assert list_paragraphs
    assert all(paragraph["style"] == "List Bullet" for paragraph in list_paragraphs)
    skills = [
        paragraph
        for paragraph in extracted["paragraphs"]
        if paragraph["content_id"].startswith("skill_groups.")
    ]
    assert all(paragraph["runs"][0]["bold"] for paragraph in skills)


def test_template_drift_is_rejected(master_resume: Path, tmp_path: Path) -> None:
    changed = tmp_path / "drifted.docx"
    shutil.copy2(master_resume, changed)
    document = Document(changed)
    document.add_paragraph("Unexpected paragraph")
    document.save(changed)
    with pytest.raises(TemplateError, match="32 paragraphs"):
        validate_template(Document(changed))
