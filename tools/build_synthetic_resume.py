#!/usr/bin/env python3
"""Build a clearly synthetic DOCX fixture for local regression tests."""

from __future__ import annotations

import argparse
import os
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.oxml import parse_xml, serialize_part_xml
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_OUTPUT = PROJECT_ROOT / "template" / "sample_resume.docx"


def _set_run_font(run: object, *, size: float = 9, bold: bool = False) -> None:
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)


def _compact(paragraph: object) -> None:
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1


def _add_bottom_rule(paragraph: object) -> None:
    properties = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "4F6B86")
    borders.append(bottom)
    properties.append(borders)


def _add_hyperlink(paragraph: object, text: str, target: str) -> None:
    relationship_id = paragraph.part.relate_to(
        target,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "315F7D")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "18")
    properties.extend((color, underline, size))
    run.append(properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _section(document: object, title: str) -> None:
    paragraph = document.add_paragraph()
    _compact(paragraph)
    run = paragraph.add_run(title)
    _set_run_font(run, size=9.5, bold=True)
    _add_bottom_rule(paragraph)


def _plain(document: object, text: str) -> None:
    paragraph = document.add_paragraph()
    _compact(paragraph)
    _set_run_font(paragraph.add_run(text))


def _labelled_bullet(document: object, label: str, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    _compact(paragraph)
    _set_run_font(paragraph.add_run(f"{label}: "), bold=True)
    _set_run_font(paragraph.add_run(text))


def _bullet(document: object, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    _compact(paragraph)
    _set_run_font(paragraph.add_run(text))


def _project_heading(document: object, name: str, technologies: str) -> None:
    paragraph = document.add_paragraph()
    _compact(paragraph)
    _set_run_font(paragraph.add_run(name), bold=True)
    _set_run_font(paragraph.add_run(" | "))
    technology_run = paragraph.add_run(technologies)
    _set_run_font(technology_run)
    technology_run.italic = True


def _strip_package_extras(output: Path) -> None:
    excluded = {
        "customXml/item1.xml",
        "customXml/_rels/item1.xml.rels",
        "customXml/itemProps1.xml",
        "docProps/thumbnail.jpeg",
    }
    with tempfile.NamedTemporaryFile(
        dir=output.parent,
        prefix=f".{output.name}.",
        suffix=".tmp",
        delete=False,
    ) as temporary:
        temporary_path = Path(temporary.name)
    try:
        with (
            zipfile.ZipFile(output, "r") as source,
            zipfile.ZipFile(temporary_path, "w", compression=zipfile.ZIP_DEFLATED) as destination,
        ):
            for info in source.infolist():
                if info.filename in excluded:
                    continue
                payload = source.read(info.filename)
                if info.filename in {"_rels/.rels", "word/_rels/document.xml.rels"}:
                    root = parse_xml(payload)
                    for relationship in list(root):
                        target = relationship.get("Target", "")
                        if "customXml/" in target or target.endswith(
                            "docProps/thumbnail.jpeg"
                        ):
                            root.remove(relationship)
                    payload = serialize_part_xml(root)
                elif info.filename == "[Content_Types].xml":
                    root = parse_xml(payload)
                    for content_type in list(root):
                        part_name = content_type.get("PartName", "")
                        extension = content_type.get("Extension", "")
                        if part_name.startswith("/customXml/") or extension == "jpeg":
                            root.remove(content_type)
                    payload = serialize_part_xml(root)
                destination.writestr(info, payload)
        os.replace(temporary_path, output)
    finally:
        temporary_path.unlink(missing_ok=True)


def build(output: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(691 / 1440)
    section.right_margin = Inches(0.5)
    section.bottom_margin = Inches(691 / 1440)
    section.left_margin = Inches(0.5)

    properties = document.core_properties
    properties.title = "Synthetic Resume Fixture"
    properties.subject = "Test fixture containing no real person or employment data"
    properties.author = "Resume Tailor synthetic tests"
    properties.last_modified_by = "Resume Tailor synthetic tests"
    properties.comments = "All names, organizations, links, and achievements are fictional."

    name = document.add_paragraph()
    _compact(name)
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(name.add_run("SAMPLE CANDIDATE"), size=14, bold=True)

    contact = document.add_paragraph()
    _compact(contact)
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    links = (
        ("Sample City", "https://example.com/location"),
        ("sample.candidate@example.com", "mailto:sample.candidate@example.com"),
        ("Portfolio", "https://example.com/portfolio"),
        ("GitHub", "https://github.com/example"),
        ("LinkedIn", "https://www.linkedin.com/in/sample-candidate"),
        ("Sample PR", "https://github.com/example/repository/pull/1"),
    )
    for index, (text, target) in enumerate(links):
        if index:
            _set_run_font(contact.add_run(" | "))
        _add_hyperlink(contact, text, target)

    _section(document, "OBJECTIVE / SUMMARY")
    _plain(
        document,
        "Synthetic engineering profile used only to test evidence-gated résumé workflows.",
    )

    _section(document, "EDUCATION & CERTIFICATIONS")
    education = document.add_paragraph()
    _compact(education)
    _set_run_font(education.add_run("Example Institute"), bold=True)
    _set_run_font(education.add_run(" | Certificate in Software Systems"))
    _labelled_bullet(document, "Coursework", "Python, testing, and data validation")
    _labelled_bullet(document, "Certifications", "Synthetic credential for fixture testing")

    _section(document, "TECHNICAL SKILLS")
    _labelled_bullet(document, "Languages", "Python, JavaScript, SQL")
    _labelled_bullet(document, "AI systems", "Structured outputs, evaluation, orchestration")
    _labelled_bullet(document, "Engineering", "FastAPI, JSON Schema, pytest, Linux")

    _section(document, "AI ENGINEERING PROJECTS")
    _project_heading(document, "Evidence Gate", "Python, JSON Schema")
    _bullet(document, "Validates synthetic claims against synthetic source evidence.")
    _bullet(document, "Stops fixture runs when required support is absent.")
    _bullet(document, "Produces deterministic validation reports for tests.")

    _project_heading(document, "Local Review Console", "FastAPI, JavaScript")
    _bullet(document, "Presents synthetic workflow stages through a localhost-only UI.")
    _bullet(document, "Requires explicit approval at each fixture review gate.")
    _bullet(document, "Escapes untrusted sample content before rendering.")
    _bullet(document, "Supports desktop and narrow-screen test layouts.")

    _project_heading(document, "Document Validator", "DOCX, LibreOffice")
    _bullet(document, "Preserves formatting in a fully synthetic one-page template.")
    _bullet(document, "Checks generated page count and required fixture text.")
    _bullet(document, "Rejects structural drift before a sample export is accepted.")

    _section(document, "OPEN SOURCE CONTRIBUTION")
    _project_heading(document, "Example Contribution", "Python, Tests")
    _bullet(document, "Represents a fictional pull request used only for hyperlink validation.")

    _section(document, "EXPERIENCE")
    experience = document.add_paragraph()
    _compact(experience)
    _set_run_font(experience.add_run("Synthetic Software Engineer"), bold=True)
    _set_run_font(experience.add_run(" | Example Organization, Sample City "))
    dates = experience.add_run("(2024–Present)")
    _set_run_font(dates)
    dates.italic = True
    _bullet(document, "Maintains fictional fixtures for safe automated workflow testing.")

    if len(document.paragraphs) != 32:
        raise RuntimeError(
            f"Synthetic fixture must contain 32 paragraphs, found {len(document.paragraphs)}."
        )
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    _strip_package_extras(output)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()
    build(args.output.resolve())
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
